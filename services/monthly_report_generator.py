# services/monthly_report_generator.py
import pandas as pd
import io
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 引用相關的資料庫查詢模組
from db import queries_salary_read as q_read
from db import queries_employee as q_emp
from db import queries_insurance as q_ins
from db import queries_attendance as q_att
from utils.helpers import get_monthly_dates

# --- 核心函式：獲取報表基礎資料 ---
def _get_monthly_salary_data(conn, year, month):
    report_df, item_types = q_read.get_salary_report_for_editing(conn, year, month)
    final_df = report_df[report_df['status'] == 'final'].copy()
    
    if final_df.empty:
        raise ValueError(f"{year} 年 {month} 月沒有任何「已鎖定」的薪資紀錄可供產生報表。")

    emp_df = q_emp.get_all_employees(conn)
    ins_df = q_ins.get_all_insurance_history(conn)
    if not ins_df.empty:
        month_start, month_end = get_monthly_dates(year, month)
        ins_df['start_date_dt'] = pd.to_datetime(ins_df['start_date'])
        ins_df['end_date_dt'] = pd.to_datetime(ins_df['end_date'], errors='coerce')

        # 1. 篩選出在該薪資月份內有效的加保紀錄
        active_ins_df = ins_df[
            (ins_df['start_date_dt'] <= pd.to_datetime(month_end)) &
            (ins_df['end_date_dt'].isnull() | (ins_df['end_date_dt'] >= pd.to_datetime(month_start)))
        ].copy()

        # 2. 在有效的紀錄中，找到每個員工最新的一筆
        if not active_ins_df.empty:
            latest_ins = active_ins_df.loc[active_ins_df.groupby('name_ch')['start_date_dt'].idxmax()]
            final_df = pd.merge(final_df, latest_ins[['name_ch', 'company_name']], left_on='員工姓名', right_on='name_ch', how='left')
        else:
            final_df['company_name'] = '無有效加保'
    else:
        final_df['company_name'] = '無加保紀錄'

    final_df = pd.merge(final_df, emp_df[['id', 'dept']], left_on='employee_id', right_on='id', how='left')
    
    for col in list(item_types.keys()) + ['應付總額', '應扣總額', '實發薪資', '匯入銀行', '現金']:
        if col in final_df.columns:
            final_df[col] = pd.to_numeric(final_df[col], errors='coerce').fillna(0)
            
    return final_df, item_types

# --- Excel 相關函式 ---
def _write_styled_excel(df: pd.DataFrame, sheet_name: str, total_cols: list) -> io.BytesIO:
    output = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    for r in dataframe_to_rows(df, index=False, header=True):
        ws.append(r)

    header_fill = PatternFill(start_color="DDEBF7", end_color="DDEBF7", fill_type="solid")
    total_fill = PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid")
    bold_font = Font(bold=True)

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = bold_font
        
    total_row = ws.max_row + 1
    ws.cell(row=total_row, column=1, value="合計").font = bold_font
    ws.cell(row=total_row, column=1).fill = total_fill

    for col_idx, col_name in enumerate(df.columns, 1):
        if col_name in total_cols:
            col_letter = ws.cell(row=1, column=col_idx).column_letter
            ws.cell(row=total_row, column=col_idx, value=f"=SUBTOTAL(9, {col_letter}2:{col_letter}{total_row-1})")
            cell_to_format = ws.cell(row=total_row, column=col_idx)
            cell_to_format.font = bold_font
            cell_to_format.fill = total_fill
            cell_to_format.number_format = '#,##0'

    for col in ws.columns:
        max_length = 0
        column_letter = col[0].column_letter
        for cell in col:
            try:
                cell_len = sum(2 if '\u4e00' <= char <= '\u9fff' else 1 for char in str(cell.value))
                if cell_len > max_length:
                    max_length = cell_len
            except:
                pass
        adjusted_width = max_length + 2
        ws.column_dimensions[column_letter].width = adjusted_width

    wb.save(output)
    output.seek(0)
    return output.getvalue()

def _generate_basic_salary_excel(conn, df: pd.DataFrame, year, month):
    att_summary = q_att.get_monthly_attendance_summary(conn, year, month)
    df_merged = pd.merge(df, att_summary, on='employee_id', how='left').fillna(0)
    
    basic_earning_items = ['底薪', '加班費(延長工時)', '加班費(再延長工時)']
    basic_deduction_items = ['勞健保', '事假', '病假', '遲到', '早退']
    all_basic_items = basic_earning_items + basic_deduction_items
    for item in all_basic_items:
        if item not in df_merged.columns:
            df_merged[item] = 0

    df_merged['應付總額'] = df_merged[basic_earning_items].sum(axis=1)
    df_merged['應扣總額'] = df_merged[basic_deduction_items].sum(axis=1)
    df_merged['實支金額'] = df_merged['應付總額'] + df_merged['應扣總額']

    cols = {
        '員工姓名': '姓名', '員工編號': '編號', 'company_name': '加保單位', 'dept': '部門', '底薪': '底薪',
        '加班費(延長工時)': '加班費', '加班費(再延長工時)': '加班費2', 
        '應付總額': '應付總額',
        '勞健保': '勞健保', '事假': '事假', '病假': '病假', 
        'late_minutes': '遲到(分)', '遲到': '遲到', 'early_leave_minutes': '早退(分)', '早退': '早退', 
        '應扣總額': '應扣總額',
        '實支金額': '實支金額',
        '勞退提撥': '勞退提撥' 
    }
    
    df_report = pd.DataFrame()
    for col_db, col_report in cols.items():
        if col_db not in df_merged.columns:
            df_merged[col_db] = 0
        df_report[col_report] = df_merged[col_db]
    
    # 定義一個安全的計算函式，用於逐行處理
    def calculate_hours(row):
        base_salary = pd.to_numeric(row.get('底薪', 0), errors='coerce')
        sija_deduction = pd.to_numeric(row.get('事假', 0), errors='coerce')
        bingjia_deduction = pd.to_numeric(row.get('病假', 0), errors='coerce')
        
        # 確保資料都成功轉為數字，否則設為 0
        base_salary = base_salary if pd.notna(base_salary) else 0
        sija_deduction = sija_deduction if pd.notna(sija_deduction) else 0
        bingjia_deduction = bingjia_deduction if pd.notna(bingjia_deduction) else 0
        
        # 計算時薪，如果底薪為0，則時薪也為0，避免除以零錯誤
        hourly_rate = base_salary / 240 if base_salary > 0 else 0
        
        # 計算時數
        sija_hours = abs(sija_deduction / hourly_rate) if hourly_rate > 0 else 0
        bingjia_hours = abs(bingjia_deduction / (hourly_rate * 0.5)) if hourly_rate > 0 else 0
        
        return pd.Series([sija_hours, bingjia_hours])

    # 使用 apply 方法將上述函式應用到每一行
    df_report[['事假(時)', '病假(時)']] = df_report.apply(calculate_hours, axis=1)

    # 最後，對結果進行四捨五入和格式化
    df_report['事假(時)'] = df_report['事假(時)'].round(2)
    df_report['病假(時)'] = df_report['病假(時)'].round(2)
    df_report['延長工時'] = (df_merged.get('overtime1_minutes', 0) / 60).round(2)
    df_report['再延長工時'] = ((df_merged.get('overtime2_minutes', 0) + df_merged.get('overtime3_minutes', 0)) / 60).round(2)

    total_cols = ['底薪', '加班費', '加班費2', '應付總額', '勞健保', '借支', '事假', '病假', '遲到', '早退', '其他', '稅款', '應扣總額', '實支金額', '勞退提撥']
    excel_output = _write_styled_excel(df_report, "薪資計算", total_cols)
    return excel_output, df_report

def _generate_full_salary_excel(df: pd.DataFrame, item_types: dict):
    df_report = df.copy()
    
    df_report.rename(columns={'實發薪資': '實支金額'}, inplace=True)
    
    df_report['加班費'] = df_report.get('加班費(延長工時)', 0) + df_report.get('加班費(再延長工時)', 0)

    # 計算申報薪資
    declaration_salary_deductions = (
        df_report.get('遲到', 0) + 
        df_report.get('早退', 0) + 
        df_report.get('事假', 0) + 
        df_report.get('病假', 0)
    )
    df_report['申報薪資'] = df_report.get('底薪', 0) + declaration_salary_deductions

    core_cols = ['員工姓名', '員工編號', 'company_name', 'dept']
    earning_cols = sorted([k for k, v in item_types.items() if v == 'earning' and '加班費' not in k and k != '底薪'])
    deduction_cols = sorted([k for k, v in item_types.items() if v == 'deduction' and k != '勞健保'])
    summary_cols = ['應付總額', '應扣總額', '實支金額', '匯入銀行', '現金', '勞退提撥']
    
    final_cols = core_cols + ['底薪', '申報薪資', '加班費'] + earning_cols + ['勞健保'] + deduction_cols + summary_cols
    
    for col in final_cols:
        if col not in df_report.columns:
            df_report[col] = 0
            
    df_report = df_report[final_cols]
    
    total_cols = [col for col in df_report.columns if col not in ['員工姓名', '員工編號', 'company_name', 'dept']]
    df_report.rename(columns={
        'company_name': '加保單位', 
        'dept': '部門',
        '勞退提撥': '勞退提撥'
    }, inplace=True)
    return _write_styled_excel(df_report, "薪資計算(加)", total_cols)

# ▼▼▼ 核心修改：已更新為您提供的最新翻譯 ▼▼▼
TRANSLATIONS = {
    # 標題
    "薪資表": {"en": "Payslip", "id": "Slip Gaji", "vi": "Phiếu Lương", "th": "สลิปเงินเดือน"},
    "姓名": {"en": "Name", "id": "Nama", "vi": "Họ và tên", "th": "ชื่อ"},
    # 表頭
    "應付項目": {"en": "Earnings", "id": "jenis Penghasilan", "vi": "Các khoản thu nhập", "th": "รายได้"},
    "金額": {"en": "Amount", "id": "Jumlah", "vi": "Số tiền", "th": "จำนวนเงิน"},
    "計算方式/單位": {"en": "Calculation / Unit", "id": "Perhitungan / Satuan", "vi": "Cách tính / Đơn vị", "th": "วิธีการคำนวณ / หน่วย"},
    "應扣項目": {"en": "Deductions", "id": "jenisPotongan", "vi": "Các khoản khấu trừ", "th": "รายการหัก"},
    # 應付項目
    "底薪": {"en": "Base Salary", "id": "Gaji Pokok", "vi": "Lương cơ bản", "th": "เงินเดือนพื้นฐาน"},
    "加班費(延長工時)": {"en": "Overtime (Regular)", "id": "Jam Lembur (Biasa)", "vi": "Tiền làm thêm giờ (Thông thường)", "th": "ค่าล่วงเวลา (ทั่วไป)"},
    "加班費(再延長工時)": {"en": "Overtime (Extended)", "id": "Jam Lembur (Lanjutan)", "vi": "Tiền làm thêm giờ (Kéo dài)", "th": "ค่าล่วงเวลา (พิเศษ)"},
    # 應扣項目
    "勞健保": {"en": "Labor & Health Insurance", "id": "asuransi tenaga kerja & asuransi kesehatan", "vi": "Bảo hiểm lao động và y tế", "th": "ประกันสังคมและสุขภาพ"},
    "事假": {"en": "Personal Leave", "id": "Cuti Pribadi", "vi": "Nghỉ việc riêng", "th": "ลากิจ"},
    "病假": {"en": "Sick Leave", "id": "Cuti Sakit", "vi": "Nghỉ ốm", "th": "ลาป่วย"},
    "遲到": {"en": "Late Arrival", "id": "Terlambat bekerja", "vi": "Đi muộn", "th": "มาสาย"},
    "早退": {"en": "Early Departure", "id": "Pulang lebih Awal", "vi": "Về sớm", "th": "กลับก่อนเวลา"},
    # 總計
    "應付合計": {"en": "Total Earnings", "id": "Total Penghasilan", "vi": "Tổng thu nhập", "th": "รายรับรวม"},
    "應扣合計": {"en": "Total Deductions", "id": "Total Potongan", "vi": "Tổng khấu trừ", "th": "ยอดหักรวม"},
    "實支金額": {"en": "Net Salary", "id": "Gaji Bersih", "vi": "Lương thực lĩnh", "th": "เงินเดือนสุทธิ"},
    "勞退提撥": {"en": "Pension Contribution", "id": "Kontribusi Pensiun", "vi": "Quỹ hưu trí", "th": "เงินสมทบกองทุนสำรองเลี้ยงชีพ"},
}

def _generate_payslip_docx(df_basic: pd.DataFrame, year: int, month: int):
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.25)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)

    style = document.styles['Normal']
    style.font.name = '標楷體'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    style.font.size = Pt(8)

    def get_multilang_text(key):
        item = TRANSLATIONS.get(key, {"en": key, "id": key, "vi": key, "th": key})
        # 1. 移除英文 'en'
        # 2. 將分隔符號從 '\n' (換行) 改成 ' / ' (空格斜線空格)
        return f"{key} / {item['id']} / {item['vi']} / {item['th']}"

    def set_cell_text(cell, text_key, is_multilang=True, bold=False, align='LEFT', size=8, is_value=False):
        p = cell.paragraphs[0]
        p.text = ""
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        if align == 'CENTER':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'RIGHT':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if is_value:
             run = p.add_run(str(text_key))
        elif is_multilang:
            run = p.add_run(get_multilang_text(text_key))
        else:
            run = p.add_run(str(text_key))

        run.font.name = '標楷體'
        run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tahoma')
        run.font.bold = bold
        run.font.size = Pt(size)

    for i, emp_row in df_basic.iterrows():
        base_salary = int(emp_row.get('底薪', 0))
        earnings = [{'label': '底薪', 'key': '底薪', 'formula': None, 'unit_key': None}, {'label': '加班費(延長工時)', 'key': '加班費', 'formula': f"{base_salary}/30/8*1.34", 'unit_key': '延長工時', 'unit': 'H'}, {'label': '加班費(再延長工時)', 'key': '加班費2', 'formula': f"{base_salary}/30/8*1.67", 'unit_key': '再延長工時', 'unit': 'H'}]
        deductions = [{'label': '勞健保', 'key': '勞健保', 'formula': None, 'unit_key': None}, {'label': '事假', 'key': '事假', 'formula': f'{base_salary}/30', 'unit_key': '事假(時)', 'unit': 'H'}, {'label': '病假', 'key': '病假', 'formula': f'{base_salary}/30/2', 'unit_key': '病假(時)', 'unit': 'H'}, {'label': '遲到', 'key': '遲到', 'formula': f'{base_salary}/30/8/60', 'unit_key': '遲到(分)', 'unit': 'M'}, {'label': '早退', 'key': '早退', 'formula': f'{base_salary}/30/8/60', 'unit_key': '早退(分)', 'unit': 'M'}]

        num_item_rows = 5
        
        table = document.add_table(rows=num_item_rows + 5, cols=6)
        table.style = 'Table Grid'
        
        widths = [Inches(2.0), Inches(0.6), Inches(1.5), Inches(2.0), Inches(0.6), Inches(1.5)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        table.cell(0, 0).merge(table.cell(0, 5))
        title_text = f"{year - 1911} 年 {month} 月份 " + TRANSLATIONS["薪資表"]["vi"] + " / " + TRANSLATIONS["薪資表"]["th"]
        set_cell_text(table.cell(0, 0), title_text, is_multilang=False, bold=True, align='CENTER', size=11)

        table.cell(1, 0).merge(table.cell(1, 5))
        name_text = get_multilang_text("姓名") + f": {emp_row.get('姓名', '')}"
        set_cell_text(table.cell(1, 0), name_text, is_multilang=False, size=10)
        
        def set_multilang_cell(cell, text_key, bold=False, align='LEFT', size=7):
            p = cell.paragraphs[0]
            p.clear()
            #p.paragraph_format.space_before = Pt(2)
            #p.paragraph_format.space_after = Pt(2)
            if align == 'CENTER': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'RIGHT': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            run = p.add_run(get_multilang_text(text_key))
            run.font.name = '標楷體'
            run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tahoma')
            run.font.bold = bold
            run.font.size = Pt(size)
        
        # 表頭
        set_multilang_cell(table.cell(2, 0), "應付項目", bold=True)
        set_multilang_cell(table.cell(2, 1), "金額", bold=True, align='RIGHT')
        set_multilang_cell(table.cell(2, 2), "計算方式/單位", bold=True, align='CENTER')
        set_multilang_cell(table.cell(2, 3), "應扣項目", bold=True)
        set_multilang_cell(table.cell(2, 4), "金額", bold=True, align='RIGHT')
        set_multilang_cell(table.cell(2, 5), "計算方式/單位", bold=True, align='CENTER')
        
        # 項目明細
        for r in range(num_item_rows):
            if r < len(earnings):
                item = earnings[r]
                amount = int(emp_row.get(item['key'], 0))
                set_multilang_cell(table.cell(r + 3, 0), item['label'])
                set_cell_text(table.cell(r + 3, 1), f"{amount:,}" if amount != 0 else "-", align='RIGHT', is_multilang=False, is_value=True, size=10)
                
                formula_text = item.get('formula', '')
                if item.get('unit_key'):
                    unit_val = emp_row.get(item['unit_key'], 0)
                    if unit_val > 0:
                        formula_text += f" ({unit_val:.2f}{item['unit']})"
                set_cell_text(table.cell(r + 3, 2), formula_text if formula_text else "-", align='CENTER', is_multilang=False, is_value=True, size=8)

            if r < len(deductions):
                item = deductions[r]
                amount = int(abs(emp_row.get(item['key'], 0)))
                set_multilang_cell(table.cell(r + 3, 3), item['label'])
                set_cell_text(table.cell(r + 3, 4), f"{amount:,}" if amount != 0 else "-", align='RIGHT', is_multilang=False, is_value=True, size=10)

                formula_text = item.get('formula', '')
                if item.get('unit_key'):
                    unit_val = emp_row.get(item['unit_key'], 0)
                    if unit_val > 0:
                        formula_text += f" ({unit_val:.2f}{item['unit']})"
                set_cell_text(table.cell(r + 3, 5), formula_text if formula_text else "-", align='CENTER', is_multilang=False, is_value=True, size=8)

        # 總計
        total_row_1 = num_item_rows + 3
        set_multilang_cell(table.cell(total_row_1, 0), '應付合計', bold=True)
        set_cell_text(table.cell(total_row_1, 1), f"{int(emp_row.get('應付總額', 0)):,}", bold=True, align='RIGHT', is_multilang=False, is_value=True, size=10)
        set_multilang_cell(table.cell(total_row_1, 3), '應扣合計', bold=True)
        set_cell_text(table.cell(total_row_1, 4), f"{int(abs(emp_row.get('應扣總額', 0))):,}", bold=True, align='RIGHT', is_multilang=False, is_value=True, size=10)
        
        total_row_2 = num_item_rows + 4
        table.cell(total_row_2, 0).merge(table.cell(total_row_2, 2))
        net_salary_text = get_multilang_text("實支金額") + f": {int(emp_row.get('實支金額', 0)):,}"
        set_cell_text(table.cell(total_row_2, 0), net_salary_text, is_multilang=False, bold=True, size=10)
        
        table.cell(total_row_2, 3).merge(table.cell(total_row_2, 5))
        pension_text = get_multilang_text("勞退提撥") + f": {int(emp_row.get('勞退提撥', 0)):,}"
        set_cell_text(table.cell(total_row_2, 3), pension_text, is_multilang=False, size=10)
        
        if (i + 1) % 3 == 0 and i < len(df_basic) - 1:
            document.add_page_break()
        # 如果不是頁面的第3筆，也同樣要判斷不是最後一筆，才插入間距
        elif i < len(df_basic) - 1:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(1) # 您可以維持或調整這個間距數字

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()

def calculate_cash_denominations(cash_payout_list: list):
    if not cash_payout_list:
        return {}
    denominations = [1000, 500, 100, 50, 10, 5, 1]
    total_counts = {d: 0 for d in denominations}
    for amount in cash_payout_list:
        remaining_amount = int(amount)
        for d in denominations:
            if remaining_amount <= 0: break
            count = remaining_amount // d
            if count > 0:
                total_counts[d] += count
                remaining_amount -= count * d
    return total_counts

def generate_monthly_salary_reports(conn, year, month):
    final_df, item_types = _get_monthly_salary_data(conn, year, month)
    
    basic_excel_file, basic_df_for_word = _generate_basic_salary_excel(conn, final_df, year, month)
    full_excel_file = _generate_full_salary_excel(final_df, item_types)
    payslip_docx_file = _generate_payslip_docx(basic_df_for_word, year, month)
    
    cash_df = final_df[final_df['現金'] > 0]
    cash_payout_list = cash_df['現金'].tolist()
    
    return {
        "basic_excel": basic_excel_file,
        "full_excel": full_excel_file,
        "payslip_docx": payslip_docx_file,
        "cash_payout_list": cash_payout_list
    }