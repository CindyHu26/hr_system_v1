# services/report_generator.py
import pandas as pd
import io
from datetime import time, datetime
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.page import PageMargins
from openpyxl.worksheet.pagebreak import Break
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from db import queries_attendance as q_att
from db import queries_salary_records as q_records
from db import queries_employee as q_emp
from db import queries_insurance as q_ins

def _write_styled_excel(df: pd.DataFrame, sheet_name: str) -> io.BytesIO:
    """將 DataFrame 寫入一個帶有格式化表頭的 Excel 檔案中。"""
    output = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    for r in dataframe_to_rows(df, index=False, header=True):
        ws.append(r)

    header_fill = PatternFill(start_color="DDEBF7", end_color="DDEBF7", fill_type="solid")
    bold_font = Font(bold=True)

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = bold_font

    for col in ws.columns:
        max_length = 0
        column_letter = col[0].column_letter
        for cell in col:
            try:
                cell_len = sum(2 if '\u4e00' <= char <= '\u9fff' else 1 for char in str(cell.value))
                if cell_len > max_length:
                    max_length = cell_len
            except:
                pass
        adjusted_width = max_length + 2
        ws.column_dimensions[column_letter].width = adjusted_width

    wb.save(output)
    output.seek(0)
    return output.getvalue()

def dataframe_to_report_excel(df_all_employees: pd.DataFrame, final_columns: list, numeric_columns: list):
    """將包含所有員工的單一 DataFrame 轉換為格式化的、可分頁列印的 Excel 檔案。"""
    output = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "出勤月報表"

    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.page_margins = PageMargins(left=0.5, right=0.5, top=0.7, bottom=0.7)
    
    ws.print_title_rows = '1:1'

    header_fill = PatternFill(start_color="CCFFCC", end_color="CCFFCC", fill_type="solid")
    bold_font = Font(bold=True)
    
    for col_num, column_title in enumerate(final_columns, 1):
        cell = ws.cell(row=1, column=col_num, value=column_title)
        cell.fill = header_fill
        cell.font = bold_font

    current_row = 2
    for name, df_employee in df_all_employees.groupby('名稱', sort=False):
        for _, row_data in df_employee.iterrows():
            for c_idx, col_name in enumerate(final_columns, 1):
                ws.cell(row=current_row, column=c_idx, value=row_data.get(col_name))
            current_row += 1
        
        subtotal_row_idx = current_row
        ws.cell(row=subtotal_row_idx, column=3, value=f"{name}_小計").font = bold_font
        for c_idx, col_name in enumerate(final_columns, 1):
            if col_name in numeric_columns:
                col_letter = ws.cell(row=1, column=c_idx).column_letter
                start_range = subtotal_row_idx - len(df_employee)
                end_range = subtotal_row_idx - 1
                ws.cell(row=subtotal_row_idx, column=c_idx, value=f"=SUM({col_letter}{start_range}:{col_letter}{end_range})").font = bold_font
        
        current_row += 1
        ws.row_breaks.append(Break(id=current_row - 1))

    column_widths = {}
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
        for cell in row:
            if cell.value:
                length = sum(2 if '\u4e00' <= char <= '\u9fff' else 1 for char in str(cell.value))
                column_widths[cell.column_letter] = max(column_widths.get(cell.column_letter, 0), length)
    
    for col, width in column_widths.items():
        if col in ['I', 'J', 'K', 'L', 'M', 'N']:
            ws.column_dimensions[col].width = 7
        else:
            ws.column_dimensions[col].width = min(max(width + 2, 8), 50)

    wb.save(output)
    output.seek(0)
    return output.getvalue()


def get_descriptive_leave_type(row):
    """根據請假時間判斷假別描述"""
    leave_type = row['leave_type']
    start_time = row['start_date'].time()
    end_time = row['end_date'].time()
    duration = row['duration']

    if pd.isna(leave_type):
        return ''
    
    if duration >= 8:
        return f"全天{leave_type}"
    if start_time < time(12, 0) and end_time <= time(13, 0):
        return f"上午{leave_type}"
    if start_time >= time(12, 0):
        return f"下午{leave_type}"
    
    return leave_type

def get_attendance_status(row):
    """根據紀錄決定出席狀態的優先級"""
    if row['缺席'] > 0:
        return '缺席'
    if row['請假'] > 0:
        return '請假/出差'
    
    is_late = row['遲到'] > 0
    is_early = row['早退'] > 0
    
    if is_late and is_early:
        return '遲到/早退'
    if is_late:
        return '遲到'
    if is_early:
        return '早退'
        
    return '一般'


def generate_attendance_excel(conn, year, month):
    """從資料庫獲取資料並產生 Excel 報表。"""
    
    df_att_raw = q_att.get_attendance_by_month(conn, year, month)
    if df_att_raw.empty:
        raise ValueError(f"{year} 年 {month} 月沒有任何出勤紀錄可供匯出。")

    df_leave_raw = q_att.get_leave_details_by_month(conn, year, month)
    
    df = df_att_raw.rename(columns={
        'hr_code': '人員 ID', 'name_ch': '名稱', 'date': '日期',
        'checkin_time': '簽到', 'checkout_time': '簽退', 'late_minutes': '遲到',
        'early_leave_minutes': '早退', 'absent_minutes': '缺席',
        'overtime1_minutes': '加班 1', 'overtime2_minutes': '加班 2',
        'overtime3_minutes': '加班 3', 'leave_minutes': '請假'
    })
    
    df['加班23'] = df['加班 2'] + df['加班 3']
    
    df['日期_dt'] = pd.to_datetime(df['日期'])
    weekday_map = {0: '一', 1: '二', 2: '三', 3: '四', 4: '五', 5: '六', 6: '日'}
    df['星期'] = df['日期_dt'].dt.weekday.map(weekday_map)
    
    if not df_leave_raw.empty:
        daily_leaves = []
        for _, leave in df_leave_raw.iterrows():
            d = leave['start_date'].date()
            while d <= leave['end_date'].date():
                daily_leaves.append({
                    'employee_id': leave['employee_id'],
                    'date_str': d.strftime('%Y-%m-%d'),
                    'leave_type': get_descriptive_leave_type(leave)
                })
                d += pd.Timedelta(days=1)
        df_daily_leaves = pd.DataFrame(daily_leaves).rename(columns={'date_str': '日期'})
        
        df = pd.merge(df, df_daily_leaves, on=['employee_id', '日期'], how='left')
        df['請假類型'] = df['leave_type'].fillna('')
    else:
        df['請假類型'] = ''
        
    df['出席狀態'] = df.apply(get_attendance_status, axis=1)
    
    df = df.sort_values(by=['人員 ID', '日期_dt'])
    
    final_columns = ['星期', '人員 ID', '名稱', '日期', '出席狀態', '請假類型', '簽到', '簽退',
                     '遲到', '早退', '缺席', '加班 1', '加班23', '請假']
    numeric_columns = ['遲到', '早退', '缺席', '加班 1', '加班23', '請假']
    
    excel_file = dataframe_to_report_excel(df, final_columns, numeric_columns)
    
    return excel_file

def _get_monthly_salary_data(conn, year, month):
    report_df, item_types = q_records.get_salary_report_for_editing(conn, year, month)
    final_df = report_df[report_df['status'] == 'final'].copy()
    
    if final_df.empty:
        raise ValueError(f"{year} 年 {month} 月沒有任何「已鎖定」的薪資紀錄可供產生報表。")

    emp_df = q_emp.get_all_employees(conn)
    ins_df = q_ins.get_all_insurance_history(conn)
    if not ins_df.empty:
        ins_df['start_date'] = pd.to_datetime(ins_df['start_date'])
        latest_ins = ins_df.loc[ins_df.groupby('name_ch')['start_date'].idxmax()]
        final_df = pd.merge(final_df, latest_ins[['name_ch', 'company_name']], left_on='員工姓名', right_on='name_ch', how='left')
    else:
        final_df['company_name'] = ''

    final_df = pd.merge(final_df, emp_df[['id', 'dept']], left_on='employee_id', right_on='id', how='left')
    
    for col in list(item_types.keys()) + ['應付總額', '應扣總額', '實發薪資']:
        if col in final_df.columns:
            final_df[col] = pd.to_numeric(final_df[col], errors='coerce').fillna(0)
            
    return final_df, item_types


def _generate_basic_salary_excel(conn, df: pd.DataFrame, year, month):
    att_summary = q_att.get_monthly_attendance_summary(conn, year, month)
    df_merged = pd.merge(df, att_summary, on='employee_id', how='left').fillna(0)

    basic_earning_items = ['底薪', '加班費(延長工時)', '加班費(再延長工時)']
    basic_deduction_items = ['勞健保', '借支', '事假', '病假', '遲到', '早退', '其他', '稅款']
    all_basic_items = basic_earning_items + basic_deduction_items
    for item in all_basic_items:
        if item not in df_merged.columns:
            df_merged[item] = 0

    df_merged['應付合計_basic'] = df_merged[basic_earning_items].sum(axis=1)
    df_merged['應扣合計_basic'] = df_merged[basic_deduction_items].sum(axis=1)
    df_merged['合計_basic'] = df_merged['應付合計_basic'] + df_merged['應扣合計_basic']

    cols = {
        '員工姓名': '姓名', '員工編號': '編號', 'company_name': '加保單位', 'dept': '部門', '底薪': '底薪',
        '加班費(延長工時)': '加班費', '加班費(再延長工時)': '加班費2', 
        '應付合計_basic': '應付合計',
        '勞健保': '勞健保', '借支': '借支', '事假': '事假', '病假': '病假', 
        'late_minutes': '遲到(分)', '遲到': '遲到', 'early_leave_minutes': '早退(分)', '早退': '早退', 
        '其他': '其他', '稅款': '稅款', 
        '應扣合計_basic': '應扣合計',
        '合計_basic': '合計',
        '勞退提撥(公司負擔)': '勞退提撥'
    }
    
    df_report = pd.DataFrame()
    for col_db, col_report in cols.items():
        if col_db not in df_merged.columns:
            df_merged[col_db] = 0
        df_report[col_report] = df_merged[col_db]
    
    hourly_rate = (df_merged['底薪'] / 240).replace(0, pd.NA)
    df_report['事假(時)'] = (df_report['事假'] / hourly_rate).abs().round(2).fillna(0)
    df_report['病假(時)'] = (df_report['病假'] / (hourly_rate * 0.5)).abs().round(2).fillna(0)
    df_report['延長工時'] = (df_merged.get('overtime1_minutes', 0) / 60).round(2)
    df_report['再延長工時'] = ((df_merged.get('overtime2_minutes', 0) + df_merged.get('overtime3_minutes', 0)) / 60).round(2)
    
    excel_output = _write_styled_excel(df_report, "薪資計算")
    return excel_output, df_report


def _generate_full_salary_excel(df: pd.DataFrame, item_types: dict):
    df_report = df.copy()
    
    df_report['加班費'] = df_report.get('加班費(延長工時)', 0) + df_report.get('加班費(再延長工時)', 0)

    core_cols = ['員工姓名', '員工編號', 'company_name', 'dept']
    earning_cols = sorted([k for k, v in item_types.items() if v == 'earning' and '加班費' not in k and k != '底薪'])
    deduction_cols = sorted([k for k, v in item_types.items() if v == 'deduction' and k != '勞健保'])
    summary_cols = ['應付總額', '應扣總額', '實發薪資', '匯入銀行', '現金', '勞退提撥(公司負擔)']
    
    final_cols = core_cols + ['底薪', '加班費'] + earning_cols + ['勞健保'] + deduction_cols + summary_cols
    
    for col in final_cols:
        if col not in df_report.columns:
            df_report[col] = 0
            
    df_report = df_report[final_cols]
    
    df_report.rename(columns={'company_name': '加保單位', 'dept': '部門'}, inplace=True)
    return _write_styled_excel(df_report, "薪資計算(加)")


def _generate_payslip_docx(df_basic: pd.DataFrame, year: int, month: int):
    """
    【最終版：左右佈局 + 完整資訊】
    產生左右結構、包含計算公式與時數的薪資單 Word 檔案。
    """
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = document.styles['Normal']
    style.font.name = '標楷體'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    style.font.size = Pt(9)

    def set_cell_text(cell, text, bold=False, align='LEFT', size=9):
        p = cell.paragraphs[0]
        p.text = ""
        p.paragraph_format.space_before = Pt(1.5)
        p.paragraph_format.space_after = Pt(1.5)
        if align == 'CENTER':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'RIGHT':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        run = p.add_run(str(text))
        run.font.name = '標楷體'
        run.font.bold = bold
        run.font.size = Pt(size)

    for i, emp_row in df_basic.iterrows():
        base_salary = int(emp_row.get('底薪', 0))

        earnings = [
            {'label': '底薪', 'key': '底薪', 'formula': None, 'unit_key': None},
            {'label': '加班費(延長工時)', 'key': '加班費', 'formula': f"{base_salary}/30/8*1.34", 'unit_key': '延長工時', 'unit': 'H'},
            {'label': '加班費(再延長工時)', 'key': '加班費2', 'formula': f"{base_salary}/30/8*1.67", 'unit_key': '再延長工時', 'unit': 'H'}
        ]
        deductions = [
            {'label': '勞健保', 'key': '勞健保', 'formula': None, 'unit_key': None},
            {'label': '借支', 'key': '借支', 'formula': None, 'unit_key': None},
            {'label': '事假', 'key': '事假', 'formula': f'{base_salary}/30', 'unit_key': '事假(時)', 'unit': 'H'},
            {'label': '病假', 'key': '病假', 'formula': f'{base_salary}/30/2', 'unit_key': '病假(時)', 'unit': 'H'},
            {'label': '遲到', 'key': '遲到', 'formula': f'{base_salary}/30/8/60', 'unit_key': '遲到(分)', 'unit': 'M'},
            {'label': '早退', 'key': '早退', 'formula': f'{base_salary}/30/8/60', 'unit_key': '早退(分)', 'unit': 'M'},
            {'label': '稅款', 'key': '稅款', 'formula': None, 'unit_key': None},
            {'label': '其他', 'key': '其他', 'formula': None, 'unit_key': None}
        ]

        num_item_rows = max(len(earnings), len(deductions))
        table = document.add_table(rows=num_item_rows + 5, cols=6)
        table.style = 'Table Grid'
        
        table.columns[0].width = Inches(1.4)
        table.columns[1].width = Inches(0.8)
        table.columns[2].width = Inches(1.45)
        table.columns[3].width = Inches(1.4)
        table.columns[4].width = Inches(0.8)
        table.columns[5].width = Inches(1.45)

        table.cell(0, 0).merge(table.cell(0, 5))
        set_cell_text(table.cell(0, 0), f"{year - 1911} 年 {month} 月份薪資表", bold=True, align='CENTER', size=12)

        table.cell(1, 0).merge(table.cell(1, 5))
        set_cell_text(table.cell(1, 0), f"姓名：{emp_row.get('姓名', '')}", size=10)

        set_cell_text(table.cell(2, 0), '應付項目', bold=True)
        set_cell_text(table.cell(2, 1), '金額', bold=True, align='RIGHT')
        set_cell_text(table.cell(2, 2), '計算方式/單位', bold=True, align='CENTER')
        set_cell_text(table.cell(2, 3), '應扣項目', bold=True)
        set_cell_text(table.cell(2, 4), '金額', bold=True, align='RIGHT')
        set_cell_text(table.cell(2, 5), '計算方式/單位', bold=True, align='CENTER')

        for r in range(num_item_rows):
            if r < len(earnings):
                item = earnings[r]
                amount = int(emp_row.get(item['key'], 0))
                set_cell_text(table.cell(r + 3, 0), item['label'])
                set_cell_text(table.cell(r + 3, 1), f"{amount:,}" if amount != 0 else "-", align='RIGHT')
                
                # --- ▼▼▼ 核心修改：將公式與單位合併為一行 ▼▼▼ ---
                formula_text = item.get('formula', '')
                if item.get('unit_key'):
                    unit_val = emp_row.get(item['unit_key'], 0)
                    if unit_val > 0:
                        formula_text += f" ({unit_val:.2f}{item['unit']})"
                set_cell_text(table.cell(r + 3, 2), formula_text if formula_text else "-", align='CENTER')

            if r < len(deductions):
                item = deductions[r]
                amount = int(abs(emp_row.get(item['key'], 0)))
                set_cell_text(table.cell(r + 3, 3), item['label'])
                set_cell_text(table.cell(r + 3, 4), f"{amount:,}" if amount != 0 else "-", align='RIGHT')

                formula_text = item.get('formula', '')
                if item.get('unit_key'):
                    unit_val = emp_row.get(item['unit_key'], 0)
                    if unit_val > 0:
                        formula_text += f" ({unit_val:.2f}{item['unit']})"
                set_cell_text(table.cell(r + 3, 5), formula_text if formula_text else "-", align='CENTER')

        total_row_1 = num_item_rows + 3
        set_cell_text(table.cell(total_row_1, 0), '應付合計', bold=True)
        set_cell_text(table.cell(total_row_1, 1), f"{int(emp_row.get('應付合計', 0)):,}", bold=True, align='RIGHT')
        set_cell_text(table.cell(total_row_1, 3), '應扣合計', bold=True)
        set_cell_text(table.cell(total_row_1, 4), f"{int(abs(emp_row.get('應扣合計', 0))):,}", bold=True, align='RIGHT')
        
        total_row_2 = num_item_rows + 4
        table.cell(total_row_2, 0).merge(table.cell(total_row_2, 2))
        set_cell_text(table.cell(total_row_2, 0), f"實支金額： {int(emp_row.get('合計', 0)):,}", bold=True)
        table.cell(total_row_2, 3).merge(table.cell(total_row_2, 5))
        set_cell_text(table.cell(total_row_2, 3), f"勞退提撥： {int(emp_row.get('勞退提撥', 0)):,}")
        
        # --- ▼▼▼ 修改：增加每頁列印筆數 ▼▼▼ ---
        if (i + 1) % 3 == 0 and i < len(df_basic) - 1:
            document.add_page_break()
        else:
            # 在非分頁處只留一個小間隔
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()

def generate_monthly_salary_reports(conn, year, month):
    """
    產生三種薪資報表（基礎Excel、完整Excel、Word薪資單）的主函式。
    """
    final_df, item_types = _get_monthly_salary_data(conn, year, month)
    
    basic_excel_file, basic_df_for_word = _generate_basic_salary_excel(conn, final_df, year, month)
    full_excel_file = _generate_full_salary_excel(final_df, item_types)
    
    payslip_docx_file = _generate_payslip_docx(basic_df_for_word, year, month)
    
    return {
        "basic_excel": basic_excel_file,
        "full_excel": full_excel_file,
        "payslip_docx": payslip_docx_file
    }